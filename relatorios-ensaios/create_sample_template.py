"""Gera modelo DOCX básico com placeholders para o sistema."""

from pathlib import Path

from docx import Document

base = Path(__file__).resolve().parent
output = base / "data" / "templates" / "modelo_exemplo.docx"
output.parent.mkdir(parents=True, exist_ok=True)

doc = Document()
doc.add_heading("RELATÓRIO TÉCNICO DE ENSAIO", level=1)

doc.add_paragraph("Relatório nº: {{ report_number }}")
doc.add_paragraph("Cliente: {{ cliente_nome }}")
doc.add_paragraph("Obra: {{ obra_nome }}")
doc.add_paragraph("Endereço: {{ obra_endereco }}")
doc.add_paragraph("Data do ensaio: {{ data_ensaio }}")
doc.add_paragraph("Responsável técnico: {{ responsavel_tecnico }}")
doc.add_paragraph("Registro profissional: {{ registro_responsavel }}")

doc.add_heading("1. Objetivo", level=2)
doc.add_paragraph("{{ objetivo }}")

doc.add_heading("2. Metodologia", level=2)
doc.add_paragraph("{{ metodologia }}")

doc.add_heading("3. Equipamentos utilizados", level=2)
doc.add_paragraph("{{ equipamentos_utilizados }}")

doc.add_heading("4. Resultados obtidos", level=2)
doc.add_paragraph("{% for item in resultados %}")
doc.add_paragraph("Item {{ item.item_number }} — {{ item.location }}")
doc.add_paragraph("Valor medido: {{ item.measured_value }} {{ item.unit }}")
doc.add_paragraph("Valor de referência: {{ item.reference_value }}")
doc.add_paragraph("Situação: {{ item.status }}")
doc.add_paragraph("Observações: {{ item.notes }}")
doc.add_paragraph("{% endfor %}")

doc.add_heading("5. Registro fotográfico", level=2)
doc.add_paragraph("{% for foto in fotos %}")
doc.add_paragraph("{{ foto.imagem }}")
doc.add_paragraph("Figura {{ loop.index }} — {{ foto.legenda }}")
doc.add_paragraph("{% endfor %}")

doc.add_heading("6. Observações técnicas", level=2)
doc.add_paragraph("{{ observacoes }}")

doc.add_heading("7. Conclusão", level=2)
doc.add_paragraph("{{ conclusao }}")

doc.save(output)
print(f"Modelo salvo em: {output}")
